import os
import json
import struct
import zlib
import pytest
from docx import Document

from generators.docx.functions import (
    generate_function_docx,
    generate_function_docx_by_file,
)
from generators.docx.appendix import generate_appendix_docx


def _minimal_png():
    """Return bytes of a valid 1x1 white PNG."""
    def chunk(chunk_type, data):
        c = chunk_type + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xff\xff\xff")
    iend = b""
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", iend)


def _get_paragraphs_text(doc):
    return [p.text for p in doc.paragraphs]


def _get_table_text(doc):
    texts = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text for cell in row.cells]
            texts.extend(row_texts)
    return texts


def _create_dummy_figures(functions, figures_dir):
    os.makedirs(figures_dir, exist_ok=True)
    png = _minimal_png()
    for func in functions:
        safe = func["name"].replace("\\", "_").replace("/", "_").replace(":", "_")
        with open(os.path.join(figures_dir, f"{safe}.png"), "wb") as f:
            f.write(png)


class TestGenerateFunctionDocx:
    def test_generates_docx_files(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_output")
        figures_dir = str(tmp_path / "figures_docx")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        for func in sample_functions:
            docx_path = os.path.join(output_dir, f"{func['name']}.docx")
            assert os.path.exists(docx_path)

    def test_docx_contains_function_name(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_output2")
        figures_dir = str(tmp_path / "figures_docx2")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "main" in all_text
        assert "Entry point" in all_text

    def test_includes_module_description_section(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_mod_desc")
        figures_dir = str(tmp_path / "figures_mod_desc")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块描述" in all_text
        assert "ModuleName<" in all_text
        assert "main" in all_text
        assert "FileName<" in all_text
        assert "main.c" in all_text
        assert "LineNumber<" in all_text
        assert "10" in all_text
        assert "MacroNameList<" in all_text
        assert "FEATURE_X" in all_text

    def test_docx_contains_input_table(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_output3")
        figures_dir = str(tmp_path / "figures_docx3")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        table_text = _get_table_text(doc)
        assert "argc" in table_text or any("argc" in p.text for p in doc.paragraphs)

    def test_empty_inputs_and_returns(self, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_empty")
        figures_dir = str(tmp_path / "figures_empty")
        funcs = [{
            "name": "empty_func",
            "file": "/proj/empty.c",
            "inputs": [],
            "returns": [],
            "body_code": "",
            "calls": [],
            "called_by": [],
            "algorithm_logic": "Does nothing.",
        }]
        _create_dummy_figures(funcs, figures_dir)

        generate_function_docx(
            function_list=funcs,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "empty_func.docx"))
        table_text = _get_table_text(doc)
        assert any("N/A" in t for t in table_text)

    def test_missing_png_does_not_crash(self, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_noimg")
        figures_dir = str(tmp_path / "figures_noimg")
        os.makedirs(figures_dir, exist_ok=True)

        funcs = [{
            "name": "no_img_func",
            "file": "/proj/empty.c",
            "inputs": [],
            "returns": [],
            "body_code": "",
            "calls": [],
            "called_by": [],
            "algorithm_logic": "",
        }]
        generate_function_docx(
            function_list=funcs,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )
        assert os.path.exists(os.path.join(output_dir, "no_img_func.docx"))

    # -- section toggle tests --

    def test_all_sections_enabled_by_default(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_default_sec")
        figures_dir = str(tmp_path / "figures_def_sec")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块描述" in all_text
        assert "输入项" in all_text
        assert "输出项" in all_text
        assert "全局数据结构" in all_text
        assert "局部数据结构" in all_text
        assert "算法和逻辑" in all_text
        assert "接口" in all_text

    def test_skip_local_data_section(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_no_local")
        figures_dir = str(tmp_path / "figures_no_local")
        _create_dummy_figures(sample_functions, figures_dir)

        sections = {"local_data": False}
        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
            sections=sections,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块描述" in all_text
        assert "输入项" in all_text
        assert "局部数据结构" not in all_text

    def test_skip_multiple_sections(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_multi_skip")
        figures_dir = str(tmp_path / "figures_multi_skip")
        _create_dummy_figures(sample_functions, figures_dir)

        sections = {"local_data": False, "algorithm": False, "interface": False}
        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
            sections=sections,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块描述" in all_text
        assert "局部数据结构" not in all_text
        assert "算法和逻辑" not in all_text
        assert "接口" not in all_text

    def test_module_summary_section_renders(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_summary")
        figures_dir = str(tmp_path / "figures_summary")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块功能" in all_text
        assert "Program entry point" in all_text

    def test_module_summary_can_be_skipped(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_no_summary")
        figures_dir = str(tmp_path / "figures_no_summary")
        _create_dummy_figures(sample_functions, figures_dir)

        sections = {"module_summary": False}
        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
            sections=sections,
        )

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "模块功能" not in all_text


class TestGroupByFileDocx:
    def test_generates_one_docx_per_source_file(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_by_file")
        figures_dir = str(tmp_path / "figures_by_file")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx_by_file(sample_functions, sample_types_json, figures_dir, output_dir)

        files = {func["file"] for func in sample_functions}
        for file_path in files:
            base = os.path.splitext(os.path.basename(file_path))[0]
            docx_path = os.path.join(output_dir, f"{base}.docx")
            assert os.path.exists(docx_path)

    def test_file_docx_contains_all_functions(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_by_file2")
        figures_dir = str(tmp_path / "figures_by_file2")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx_by_file(sample_functions, sample_types_json, figures_dir, output_dir)

        doc = Document(os.path.join(output_dir, "main.docx"))
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "main" in all_text
        assert "init" in all_text
        assert "process" in all_text

    def test_group_by_file_via_dispatcher(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_group_file")
        figures_dir = str(tmp_path / "figures_group")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
            group_by="file",
        )

        assert os.path.exists(os.path.join(output_dir, "main.docx"))

    def test_default_group_by_is_function(self, sample_functions, sample_types_json, tmp_path):
        output_dir = str(tmp_path / "docx_default")
        figures_dir = str(tmp_path / "figures_default")
        _create_dummy_figures(sample_functions, figures_dir)

        generate_function_docx(
            function_list=sample_functions,
            types_json=sample_types_json,
            figures_dir=figures_dir,
            output_dir=output_dir,
        )

        assert os.path.exists(os.path.join(output_dir, "main.docx"))
        assert os.path.exists(os.path.join(output_dir, "init.docx"))
        assert os.path.exists(os.path.join(output_dir, "process.docx"))


class TestGenerateAppendixDocx:
    def test_generates_appendix(self, sample_types_json, tmp_path):
        output_path = str(tmp_path / "appendix.docx")
        generate_appendix_docx(sample_types_json, output_path)
        assert os.path.exists(output_path)

        doc = Document(output_path)
        all_text = " ".join(_get_paragraphs_text(doc))
        assert "Appendix Global Data Structures" in all_text

    def test_contains_all_types(self, sample_types_json, tmp_path):
        output_path = str(tmp_path / "appendix2.docx")
        generate_appendix_docx(sample_types_json, output_path)

        doc = Document(output_path)
        table_text = _get_table_text(doc)
        joined = " ".join(table_text)
        assert "Point" in joined
        assert "Direction" in joined
        assert "BYTE" in joined

    def test_uses_a_reference_labels(self, sample_types_json, tmp_path):
        output_path = str(tmp_path / "appendix3.docx")
        generate_appendix_docx(sample_types_json, output_path)

        doc = Document(output_path)
        table_text = _get_table_text(doc)
        joined = " ".join(table_text)
        assert "A_1" in joined
        assert "A_2" in joined
        assert "A_3" in joined

    def test_missing_types_file(self, tmp_path):
        output_path = str(tmp_path / "appendix_missing.docx")
        generate_appendix_docx("/nonexistent/path.json", output_path)
        # Should not crash, just log warning

    def test_type_description(self, sample_types_json, tmp_path):
        output_path = str(tmp_path / "appendix_desc.docx")
        generate_appendix_docx(sample_types_json, output_path)

        doc = Document(output_path)
        table_text = _get_table_text(doc)
        joined = " ".join(table_text)
        assert "2D coordinate point" in joined
